import httpx
import json
import re

from docx import Document
from docx.shared import Pt, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import DEEPSEEK_BASE_URL

_SYSTEM_PROMPT = """Ты — опытный IT-рекрутер, который составляет реалистичные резюме на основе описания кандидата.

ГЛАВНЫЕ ПРАВИЛА — соблюдай их строго:

1. КАРЬЕРНЫЙ РОСТ — роли в проектах должны отражать логичное развитие карьеры:
   - Ранние проекты (последние в списке, самые старые) = junior/middle роли, небольшие команды
   - Средние проекты = middle/senior, растущая ответственность
   - Свежие проекты (первые в списке) = senior/lead/architect/principal engineer
   - Типичная прогрессия: Junior Developer → Developer → Senior Developer → Tech Lead → Principal Engineer / Engineering Manager
   - НЕ делай все проекты с одинаковой ролью — каждый шаг карьеры должен быть виден

2. СУММА ДЛИТЕЛЬНОСТЕЙ = ОБЩИЙ ОПЫТ:
   - Сложи месяцы всех проектов — сумма должна совпадать с полем "experience"
   - Пример: experience "8 лет" = проекты суммарно ~96 месяцев
   - Допускается ±6 месяцев (перекрытия, небольшие паузы между работами)
   - Каждый проект: минимум 6 месяцев, максимум 60 месяцев (реалистично)

3. КОЛИЧЕСТВО ПРОЕКТОВ — строго определяется пользователем в промпте:
   - Если в промпте указано "4 проекта" — генерируй ровно 4
   - Если количество не указано — определи сам исходя из опыта

4. ТЕХНОЛОГИИ РАЗВИВАЮТСЯ:
   - В старых проектах — более старые/базовые стеки
   - В новых проектах — современные технологии
   - Не вставляй Kubernetes или GraphQL в проект 2008 года

5. РАЗМЕР КОМАНДЫ РАСТЁТ С КАРЬЕРОЙ:
   - Junior-проект: "3-5 человек", "небольшая команда"
   - Senior/Lead: "15-20 Backend, 8 Frontend, 3 QA"
   - Principal/Architect: "40+ Backend, 15 Frontend, 8 DevOps, 5 QA"

6. РЕАЛИЗМ ОПИСАНИЙ:
   - Описание проекта (поле "description") — ровно 2-3 предложения о самом продукте/системе:
     * 1-е предложение: что представляет собой проект (какой продукт, платформа, система)
     * 2-е предложение: функционал и возможности системы, её составные части
     * 3-е предложение: масштаб, аудитория, технические требования к системе
     * Пример: "Проект представляет собой высоконагруженную платформу для онлайн-торговли электроникой и бытовой техникой. Система включает каталог товаров, модуль рекомендаций, платёжный шлюз и личный кабинет покупателя. Платформа обслуживает более 2 млн уникальных пользователей в сутки с требованием доступности 99.9%."
   - "Что реализовывал" — конкретные технические задачи, не абстрактные глаголы
   - Плохо: "Разрабатывал backend"
   - Хорошо: "Спроектировал микросервисную архитектуру для обработки 50k RPS на пиковой нагрузке"
   - 5-8 конкретных пунктов в implementation

Верни ТОЛЬКО валидный JSON без markdown-обёртки:
{
  "name": "Имя Фамилия",
  "specialization": "Актуальная должность (последняя/текущая)",
  "experience": "X лет Y месяцев",
  "languages": "язык программирования 1, язык программирования 2 (ТОЛЬКО языки программирования, не человеческие языки)",
  "frameworks": "фреймворк1, фреймворк2",
  "libraries": "либа1, либа2",
  "other_skills": "Docker, Linux, ...",
  "projects": [
    {
      "name": "Название компании",
      "role": "Роль на ЭТОМ конкретном проекте",
      "team": "Конкретный состав: X Backend, Y Frontend, Z QA",
      "duration": "XX месяцев",
      "description": "2-3 предложения о самом проекте/продукте: что это, его функционал, масштаб и аудитория",
      "implementation": ["конкретная задача 1", "конкретная задача 2"],
      "tech_stack": "конкретные технологии этого проекта"
    }
  ]
}
Проекты сортируй от НОВЕЙШЕГО к СТАРЕЙШЕМУ (первый = текущее/последнее место работы).
Никакого текста вне JSON. Только JSON."""

_REGEN_FIELD_HINTS = {
    "name": "Предложи имя и фамилию IT-специалиста.",
    "specialization": "Предложи должность/специализацию.",
    "experience": "Укажи суммарный опыт в формате 'X лет Y месяцев'.",
    "languages": "Перечисли языки программирования через запятую.",
    "frameworks": "Перечисли фреймворки через запятую.",
    "libraries": "Перечисли библиотеки и СУБД через запятую.",
    "other_skills": "Перечисли прочие технические навыки через запятую.",
    "description": "Опиши сам проект/продукт в 2-3 предложениях: 1) что представляет собой система/платформа, 2) её функционал и составные части, 3) масштаб, аудитория и технические требования.",
    "role": "Укажи роль специалиста в этом проекте.",
    "team": "Опиши состав и размер команды.",
    "duration": "Укажи длительность работы в формате 'X месяцев'.",
    "implementation": "Перечисли конкретные задачи, которые реализовал специалист (JSON-массив строк).",
    "tech_stack": "Перечисли технологии, используемые в проекте, через запятую.",
}


def _parse_json(raw: str) -> dict | list:
    raw = raw.strip()
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


async def generate_cv_data(prompt: str, api_key: str, model: str) -> dict:
    async with httpx.AsyncClient(timeout=120.0) as client:
        resp = await client.post(
            f"{DEEPSEEK_BASE_URL}/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": model,
                "messages": [
                    {"role": "system", "content": _SYSTEM_PROMPT},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.7,
                "max_tokens": 4096,
                "response_format": {"type": "json_object"},
            },
        )
        resp.raise_for_status()
        data = resp.json()
        choice = data["choices"][0]
        # Warn if model stopped mid-generation
        finish_reason = choice.get("finish_reason", "")
        if finish_reason == "length":
            raise ValueError(
                "Ответ модели обрезан (превышен лимит токенов). "
                "Попробуйте уменьшить количество проектов или упростить промпт."
            )
        content = choice["message"]["content"]
        return _parse_json(content)


async def regen_field(
    field: str,
    context: dict,
    hint: str,
    api_key: str,
    model: str,
) -> str | list:
    ctx_lines = "\n".join(f"{k}: {v}" for k, v in context.items() if v)
    system = (
        f"Ты — специалист по IT-резюме. Верни ТОЛЬКО значение поля «{field}» без пояснений.\n"
        f"Для поля «implementation» верни JSON-массив строк.\n"
        f"Для остальных полей верни обычную строку.\n\n"
        f"Контекст CV:\n{ctx_lines}"
    )
    user_msg = _REGEN_FIELD_HINTS.get(field, f"Сгенерируй значение для поля {field}.")
    if hint:
        user_msg += f"\nДополнительные требования: {hint}"

    async with httpx.AsyncClient(timeout=45.0) as client:
        resp = await client.post(
            f"{DEEPSEEK_BASE_URL}/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": model,
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user_msg},
                ],
                "temperature": 0.7,
            },
        )
        resp.raise_for_status()
        result = resp.json()["choices"][0]["message"]["content"].strip()

    if field == "implementation":
        try:
            parsed = _parse_json(result)
            if isinstance(parsed, list):
                return parsed
        except Exception:
            pass
        return [
            line.strip().lstrip("•–-").strip()
            for line in result.split("\n")
            if line.strip()
        ]
    return result


# ── DOCX Generation ────────────────────────────────────────────────────────

def _set_table_width(table, width_dxa: int):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _set_cell_borders(cell, color: str = "BFBFBF"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _add_skill_row(table, label: str, value: str):
    row = table.add_row()
    cell = row.cells[0]
    _set_cell_borders(cell)
    para = cell.paragraphs[0]
    r_label = para.add_run(f"{label}: ")
    r_label.font.name = "Arial"
    r_label.font.bold = True
    r_label.font.size = Pt(11)
    r_value = para.add_run(value or "")
    r_value.font.name = "Nunito"
    r_value.font.bold = False
    r_value.font.size = Pt(11)


def generate_docx(cv_data: dict, output_path: str, logo_path: str):
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Header — logo
    header = section.header
    h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    h_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h_para.clear()
    run = h_para.add_run()
    run.add_picture(logo_path, width=Emu(1368701), height=Emu(190500))

    # Title: Name — Specialization
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title_para.add_run(f"{cv_data['name']} — {cv_data['specialization']}")
    r.font.name = "Arial"
    r.font.bold = True
    r.font.size = Pt(13)

    # Skills table
    skills_table = doc.add_table(rows=0, cols=1)
    _set_table_width(skills_table, 10256)
    for label, key in [
        ("Опыт", "experience"),
        ("Языки программирования", "languages"),
        ("Фреймворки", "frameworks"),
        ("Библиотеки", "libraries"),
        ("Также опыт", "other_skills"),
    ]:
        _add_skill_row(skills_table, label, cv_data.get(key, ""))

    # Section heading
    doc.add_paragraph()
    h2 = doc.add_paragraph()
    r2 = h2.add_run("Ключевые проекты")
    r2.font.name = "Arial"
    r2.font.bold = True
    r2.font.size = Pt(12)

    # Per-project tables
    for project in cv_data.get("projects", []):
        # Project name subheading
        pn = doc.add_paragraph()
        rpn = pn.add_run(project.get("name", ""))
        rpn.font.name = "Arial"
        rpn.font.bold = True
        rpn.font.size = Pt(11)

        pt = doc.add_table(rows=0, cols=1)
        _set_table_width(pt, 10256)

        _add_skill_row(pt, "Описание проекта", project.get("description", ""))
        _add_skill_row(pt, "Роль", project.get("role", ""))
        _add_skill_row(pt, "Команда", project.get("team", ""))

        # Implementation — bullets
        impl_row = pt.add_row()
        impl_cell = impl_row.cells[0]
        _set_cell_borders(impl_cell)
        ip = impl_cell.paragraphs[0]
        r_il = ip.add_run("Что реализовывал: ")
        r_il.font.name = "Arial"
        r_il.font.bold = True
        r_il.font.size = Pt(11)
        impl = project.get("implementation", [])
        impl_text = (
            "\n".join(f"• {i}" for i in impl)
            if isinstance(impl, list)
            else str(impl)
        )
        r_iv = ip.add_run(impl_text)
        r_iv.font.name = "Nunito"
        r_iv.font.size = Pt(11)

        _add_skill_row(pt, "Стек", project.get("tech_stack", ""))
        _add_skill_row(pt, "Длительность", project.get("duration", ""))

        doc.add_paragraph()  # spacer between projects

    doc.save(output_path)
